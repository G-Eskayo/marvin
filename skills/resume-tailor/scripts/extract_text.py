#!/usr/bin/env python3
"""
extract_text.py — extract plain text from PDF, DOCX, MD, or TXT files.
Usage: python extract_text.py <filepath>
Prints plain text to stdout. Exits non-zero on error.
"""
import sys
from pathlib import Path


def extract(path: str) -> str:
    p = Path(path).expanduser().resolve()
    if not p.exists():
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        sys.exit(1)

    suffix = p.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(p)
    elif suffix == ".docx":
        return _extract_docx(p)
    elif suffix in (".md", ".txt"):
        return p.read_text(encoding="utf-8", errors="replace")
    else:
        print(f"ERROR: unsupported file type '{suffix}'. Supported: .pdf .docx .md .txt", file=sys.stderr)
        sys.exit(1)


def _extract_pdf(p: Path) -> str:
    try:
        from pdfminer.high_level import extract_text as pdf_extract
    except ImportError:
        print("ERROR: missing dependency. Run: ~/.agents/venv/bin/pip install pdfminer.six", file=sys.stderr)
        sys.exit(2)

    text = pdf_extract(str(p))
    if not text or not text.strip():
        print(f"ERROR: no text extracted from {p.name} — may be a scanned/image PDF", file=sys.stderr)
        sys.exit(1)
    return text.strip()


def _extract_docx(p: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        print("ERROR: missing dependency. Run: ~/.agents/venv/bin/pip install python-docx", file=sys.stderr)
        sys.exit(2)

    doc = Document(str(p))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: extract_text.py <filepath>", file=sys.stderr)
        sys.exit(2)
    print(extract(sys.argv[1]))
